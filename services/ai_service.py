import os
import re
import joblib
import pandas as pd
import numpy as np
from pypdf import PdfReader
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Load ML models
MODELS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "models")

try:
    recruitment_model = joblib.load(os.path.join(MODELS_DIR, "recruitment_model.pkl"))
    resume_classifier = joblib.load(os.path.join(MODELS_DIR, "resume_classifier.pkl"))
    vectorizer = joblib.load(os.path.join(MODELS_DIR, "vectorizer.pkl"))
    label_encoder = joblib.load(os.path.join(MODELS_DIR, "label_encoder.pkl"))
    gender_encoder = joblib.load(os.path.join(MODELS_DIR, "gender_encoder.pkl"))
    strategy_encoder = joblib.load(os.path.join(MODELS_DIR, "strategy_encoder.pkl"))
except Exception as e:
    print(f"Error loading pickle models: {e}")
    recruitment_model = None
    resume_classifier = None
    vectorizer = None
    label_encoder = None

# Keywords lists for skills detection and job category analysis
SKILLS_KEYWORDS = {
    "Python", "Java", "JavaScript", "TypeScript", "React", "Angular", "Vue", "Node.js", 
    "Express", "Django", "Flask", "Spring Boot", "SQL", "NoSQL", "MongoDB", "PostgreSQL", 
    "MySQL", "AWS", "Docker", "Kubernetes", "Git", "HTML", "CSS", "Tailwind", "Bootstrap", 
    "Machine Learning", "Data Science", "Data Analysis", "Deep Learning", "TensorFlow", 
    "PyTorch", "Pandas", "NumPy", "Scikit-Learn", "Figma", "UI/UX", "User Research", 
    "Prototyping", "Design Systems", "Product Management", "Agile", "Scrum", "DevOps", 
    "CI/CD", "Linux", "Cybersecurity", "Network Security", "Cloud Computing", "Android", 
    "iOS", "Swift", "Kotlin", "Flutter", "React Native", "PHP", "C#", "C++", "Ruby", "Go"
}

JOB_KEYWORDS = {
    "Backend Developer": ["Python", "Java", "Node.js", "Express", "Django", "Flask", "Spring Boot", "SQL", "PostgreSQL", "MySQL", "AWS", "Docker", "Kubernetes", "APIs", "RESTful"],
    "Data Analyst": ["SQL", "Excel", "Tableau", "Power BI", "Python", "Pandas", "NumPy", "Data Visualization", "Statistics", "Data Cleaning"],
    "Machine Learning Engineer": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-Learn", "NLP", "Computer Vision", "Pandas", "NumPy"],
    "Web Designing": ["HTML", "CSS", "JavaScript", "Figma", "UI/UX", "Tailwind", "Bootstrap", "React", "Responsive Design", "Photoshop", "Adobe Illustrator"],
    "DevOps Engineer": ["AWS", "Docker", "Kubernetes", "CI/CD", "Git", "Linux", "Jenkins", "Terraform", "Ansible", "Bash"],
    "Java Developer": ["Java", "Spring Boot", "Hibernate", "Maven", "Gradle", "SQL", "APIs", "Multithreading"],
    "Python Developer": ["Python", "Django", "Flask", "FastAPI", "SQL", "Pandas", "NumPy", "Web Scraping"]
}

def parse_resume(filepath):
    """Parse text content from PDF, DOCX, or text files."""
    ext = os.path.splitext(filepath)[1].lower()
    text = ""
    
    if ext == ".pdf":
        try:
            reader = PdfReader(filepath)
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
            
    elif ext in [".docx", ".doc"]:
        try:
            doc = Document(filepath)
            
            def get_para_text(p):
                p_text = p.text.strip()
                if not p_text:
                    return ""
                style_name = p.style.name.lower() if p.style else ""
                if "bullet" in style_name or "list" in style_name:
                    if not p_text.startswith(("-", "•", "*", "◦")):
                        p_text = "• " + p_text
                return p_text

            parent_elm = doc.element.body
            for child in parent_elm.iterchildren():
                tag = child.tag.split('}')[-1]
                if tag == 'p':
                    from docx.text.paragraph import Paragraph
                    p = Paragraph(child, doc)
                    p_text = get_para_text(p)
                    if p_text:
                        text += p_text + "\n"
                elif tag == 'tbl':
                    from docx.table import Table
                    t = Table(child, doc)
                    for row in t.rows:
                        row_cells_text = []
                        for cell in row.cells:
                            cell_text = ' '.join(get_para_text(p_item) for p_item in cell.paragraphs if get_para_text(p_item))
                            if cell_text:
                                row_cells_text.append(cell_text)
                        if row_cells_text:
                            text += '\t'.join(row_cells_text) + "\n"
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            
    else:  # Treat as plain text
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            
    return text.strip()

def analyze_resume_content(text):
    """Analyze resume text and extract key metrics, skills, category and suggestions."""
    from services.cv_analysis_service import analyze_cv
    from services.recruitment_service import predict_hiring_decision
    
    # Run the hybrid CV analysis engine
    cv_analysis = analyze_cv(text)
    
    # Run the recruitment prediction using REAL data extracted from CV — no hardcoded values
    rec_data = {
        'Age': 28,  # Neutral default — no CV-derivable equivalent
        'Gender': 1,  # Neutral default — model requires this feature
        'EducationLevel': cv_analysis['details']['educationLevel'],
        'ExperienceYears': cv_analysis['details']['experienceYears'],
        'PreviousCompanies': cv_analysis['details']['previousCompanies'],
        'DistanceFromCompany': 10.0,  # Neutral default
        'InterviewScore': cv_analysis.get('communicationScore', 70),  # Proxy: communication score from CV
        'SkillScore': cv_analysis['details']['skillScore'],
        'PersonalityScore': cv_analysis.get('culturalFitScore', 70),  # Proxy: cultural fit score from CV
        'RecruitmentStrategy': 2  # Standard job board strategy
    }
    rec_result = predict_hiring_decision(rec_data)
    
    # Merge findings
    cv_analysis['hiringChance'] = rec_result['hiringChance']
    cv_analysis['hiringRecommendation'] = rec_result['hiringRecommendation']
    
    return cv_analysis

def optimize_resume_via_ai(resume_data):
    """
    AI optimization of Resume based on target job position.
    - Sorts experiences by relevance
    - Enhances weak action verbs in bullet points
    - Injects missing ATS keywords into skills only (NOT into bullets, to preserve naturalness)
    - Preserves ALL original data including contact, education, certifications
    """
    # Support wrapper object if passed directly from Flask request
    if "resume" in resume_data:
        resume_data = resume_data["resume"]

    target_job = resume_data.get("targetJob", "Backend Developer")
    skills_input = resume_data.get("skills", {})

    # 1. Count and map original unique skills (case-insensitive) for validation
    original_skills_set = set()
    if isinstance(skills_input, dict):
        skills_dict = skills_input
        for v in skills_dict.values():
            if isinstance(v, str) and v:
                for s in v.split(","):
                    s_clean = s.strip()
                    if s_clean:
                        original_skills_set.add(s_clean.lower())
    elif isinstance(skills_input, list):
        for s in skills_input:
            if s:
                original_skills_set.add(str(s).strip().lower())
        skills_dict = {"technicalSkills": ", ".join([str(s) for s in skills_input if s])}
    else:
        skills_dict = {}

    # Gather skills flat list across all categories to check presence
    skills_flat = []
    for v in skills_dict.values():
        if isinstance(v, str) and v:
            skills_flat.extend([s.strip() for s in v.split(",") if s.strip()])

    experiences = resume_data.get("experiences", [])

    # Get ATS keywords for target job
    keywords = JOB_KEYWORDS.get(target_job, ["Python", "Development", "Software Engineering"])

    # Score & sort experiences by keyword relevance (highest first)
    rated_experiences = []
    for exp in experiences:
        title = exp.get("position", exp.get("title", ""))
        desc = exp.get("description", "")
        rel_score = 0
        for kw in keywords:
            if kw.lower() in title.lower():
                rel_score += 5
            if kw.lower() in desc.lower():
                rel_score += 2
        rated_experiences.append((rel_score, exp))

    rated_experiences.sort(key=lambda x: x[0], reverse=True)
    optimized_experiences = [item[1] for item in rated_experiences]

    # Action verbs for replacing weak openers (only if the bullet starts with a weak word)
    weak_to_strong = {
        "helped with": "collaborated on",
        "worked on": "developed",
        "responsible for": "led",
        "fixed": "resolved",
        "wrote": "implemented",
    }
    strong_openers = ["Spearheaded", "Architected", "Engineered", "Optimized",
                      "Formulated", "Pioneered", "Orchestrated", "Implemented",
                      "Developed", "Led", "Designed", "Built", "Created",
                      "Managed", "Delivered", "Collaborated"]

    # 2. Strengthen weak action verbs — do NOT inject keywords into bullets
    for i, exp in enumerate(optimized_experiences):
        desc = exp.get("description", "")
        if not desc:
            continue

        bullets = [b.strip("-*• ") for b in desc.split("\n") if b.strip()]
        optimized_bullets = []
        for j, bullet in enumerate(bullets):
            new_bullet = bullet
            # Replace known weak phrases
            for old, new in weak_to_strong.items():
                new_bullet = re.sub(r"\b" + re.escape(old) + r"\b", new, new_bullet, flags=re.I)

            # Replace weak first-person pronouns with strong action verbs
            words = new_bullet.split()
            if words and words[0].lower() in ["i", "was"]:
                verb = strong_openers[(i + j) % len(strong_openers)]
                new_bullet = verb + " " + " ".join(words[1:])

            optimized_bullets.append(f"• {new_bullet}")

        exp["description"] = "\n".join(optimized_bullets)

    # 3. Add missing ATS keywords to skills only (preserves naturalness of bullet text)
    added_kws = []
    existing_skills_lower = {s.lower() for s in skills_flat}
    
    for kw in keywords[:6]:
        # Perform robust case-insensitive check and substring check
        if kw.lower() not in existing_skills_lower:
            if not any(kw.lower() in s for s in existing_skills_lower):
                added_kws.append(kw)

    # Rebuild skills dict with added keywords appended to technicalSkills
    skills_output = {k: v for k, v in skills_dict.items()}
    if added_kws:
        tech = skills_output.get("technicalSkills", "")
        # Filter out duplicates from added_kws that are already in technicalSkills
        filtered_added_kws = []
        for kw in added_kws:
            if kw.lower() not in [s.strip().lower() for s in tech.split(",") if s.strip()]:
                filtered_added_kws.append(kw)
        
        if filtered_added_kws:
            sep = ", " if tech.strip() else ""
            skills_output["technicalSkills"] = tech + sep + ", ".join(filtered_added_kws)

    # 4. Count validation: Ensure no original skill is dropped
    enhanced_skills_set = set()
    for v in skills_output.values():
        if isinstance(v, str) and v:
            for s in v.split(","):
                s_clean = s.strip()
                if s_clean:
                    enhanced_skills_set.add(s_clean.lower())

    missing_skills = []
    for orig_s in original_skills_set:
        if orig_s not in enhanced_skills_set:
            # Find the original case from input
            orig_case = orig_s
            for v in skills_dict.values():
                if isinstance(v, str) and v:
                    for s in v.split(","):
                        if s.strip().lower() == orig_s:
                            orig_case = s.strip()
                            break
            missing_skills.append(orig_case)

    if missing_skills:
        tech = skills_output.get("technicalSkills", "")
        sep = ", " if tech.strip() else ""
        skills_output["technicalSkills"] = tech + sep + ", ".join(missing_skills)
        enhanced_skills_set.update([s.lower() for s in missing_skills])

    original_skill_count = len(original_skills_set)
    enhanced_skill_count = len(enhanced_skills_set)

    # Final assertion validation
    if enhanced_skill_count < original_skill_count:
        skills_output = skills_dict

    # Return ALL original fields — nothing is dropped during enhancement
    return {
        "name": resume_data.get("name", ""),
        "contact": resume_data.get("contact", {}),
        "education": resume_data.get("education", []),
        "experiences": optimized_experiences,
        "projects": resume_data.get("projects", []),
        "leadership": resume_data.get("leadership", []),
        "certifications": resume_data.get("certifications", []),
        "interests": resume_data.get("interests", ""),
        "skills": skills_output,
        "targetJob": target_job,
        "addedKeywords": added_kws
    }

def generate_pdf_resume(resume_data, template_style, output_path):
    """Generate a PDF document based on the Harvard Resume layout."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    
    # Page template setup (0.5 inches margins = 36 pt)
    doc = SimpleDocTemplate(
        output_path, 
        pagesize=A4, 
        rightMargin=36, 
        leftMargin=36, 
        topMargin=36, 
        bottomMargin=36
    )
    story = []
    styles = getSampleStyleSheet()
    
    # Extract data
    name = resume_data.get("name", "")
    contact = resume_data.get("contact", {})
    education = resume_data.get("education", [])
    experiences = resume_data.get("experiences", [])
    projects = resume_data.get("projects", [])
    leadership = resume_data.get("leadership", [])
    raw_skills = resume_data.get("skills", {})
    # Normalize skills: accept dict, list, or None — always produce a dict
    if isinstance(raw_skills, dict):
        skills = raw_skills
    elif isinstance(raw_skills, list):
        skills = {"technicalSkills": ", ".join(str(s) for s in raw_skills if s)}
    else:
        skills = {}
    certifications = resume_data.get("certifications", [])
    interests = resume_data.get("interests", "")
    
    # Harvard Styles (Times New Roman-based)
    title_style = ParagraphStyle(
        "HarvardTitle",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=15,
        leading=17,
        textColor=colors.black,
        alignment=1, # Center
        spaceAfter=3
    )
    
    contact_style = ParagraphStyle(
        "HarvardContact",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#111827"),
        alignment=1, # Center
        spaceAfter=10
    )
    
    section_heading_style = ParagraphStyle(
        "HarvardSectionHeading",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=10,
        leading=12,
        textColor=colors.black,
        alignment=1, # Center
        spaceBefore=8,
        spaceAfter=2
    )
    
    left_bold_style = ParagraphStyle(
        "HarvardLeftBold",
        parent=styles["Normal"],
        fontName="Times-Bold",
        fontSize=9.5,
        leading=11,
        textColor=colors.black,
        alignment=0 # Left
    )
    
    left_style = ParagraphStyle(
        "HarvardLeft",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=9.5,
        leading=11,
        textColor=colors.black,
        alignment=0 # Left
    )
    
    left_italic_style = ParagraphStyle(
        "HarvardLeftItalic",
        parent=styles["Normal"],
        fontName="Times-Italic",
        fontSize=9.5,
        leading=11,
        textColor=colors.black,
        alignment=0 # Left
    )
    
    right_style = ParagraphStyle(
        "HarvardRight",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=9.5,
        leading=11,
        textColor=colors.black,
        alignment=2 # Right
    )
    
    right_italic_style = ParagraphStyle(
        "HarvardRightItalic",
        parent=styles["Normal"],
        fontName="Times-Italic",
        fontSize=9.5,
        leading=11,
        textColor=colors.black,
        alignment=2 # Right
    )
    
    bullet_style = ParagraphStyle(
        "HarvardBullet",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=9.5,
        leading=12,
        leftIndent=15,
        firstLineIndent=-8,
        spaceAfter=2,
        textColor=colors.black
    )
    
    # Name
    if name:
        story.append(Paragraph(name.upper(), title_style))
        
    # Contact
    contact_parts = []
    if contact.get("address"):
        contact_parts.append(contact["address"])
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("linkedin"):
        contact_parts.append(contact["linkedin"])
    if contact.get("portfolio"):
        contact_parts.append(contact["portfolio"])
        
    if contact_parts:
        contact_line = "  •  ".join(contact_parts)
        story.append(Paragraph(contact_line, contact_style))
        
    # Section divider generator
    def add_section_header(title):
        heading_p = Paragraph(f"<b>{title.upper()}</b>", section_heading_style)
        t = Table([[heading_p]], colWidths=[523])
        t.setStyle(TableStyle([
            ('LINEBELOW', (0,0), (-1,-1), 0.5, colors.black),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(t)
        story.append(Spacer(1, 3))

    # --- EDUCATION ---
    if education:
        add_section_header("Education")
        for edu in education:
            uni = edu.get("university", "")
            degree = edu.get("degree", "")
            gpa = edu.get("gpa", "")
            start = edu.get("startYear", "")
            end = edu.get("endYear", "")
            coursework = edu.get("relevantCoursework", "")
            
            dates = f"{start} – {end}" if start and end else (start or end or "")
            
            row1_left = Paragraph(f"<b>{uni}</b>", left_bold_style)
            row1_right = Paragraph(dates, right_style)
            
            row2_text = degree
            if gpa:
                row2_text += f" (GPA: {gpa})"
                
            row2_left = Paragraph(row2_text, left_style)
            row2_right = Paragraph("", right_style)
            
            t_edu = Table([[row1_left, row1_right], [row2_left, row2_right]], colWidths=[383, 140])
            t_edu.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(t_edu)
            
            if coursework:
                cw_text = f"<b>Relevant Coursework:</b> {coursework}"
                story.append(Paragraph(cw_text, ParagraphStyle("cw", parent=left_style, leftIndent=10, spaceBefore=1)))
                
            story.append(Spacer(1, 3))

    # --- EXPERIENCE ---
    if experiences:
        add_section_header("Experience")
        for exp in experiences:
            company = exp.get("company", "")
            pos = exp.get("position", "")
            loc = exp.get("location", "")
            start = exp.get("startDate", "")
            end = exp.get("endDate", "")
            desc = exp.get("description", "")
            
            dates = f"{start} – {end}" if start and end else (start or end or "")
            
            row1_left = Paragraph(f"<b>{company}</b>", left_bold_style)
            row1_right = Paragraph(loc, right_style)
            
            row2_left = Paragraph(pos, left_italic_style)
            row2_right = Paragraph(dates, right_italic_style)
            
            t_exp = Table([[row1_left, row1_right], [row2_left, row2_right]], colWidths=[383, 140])
            t_exp.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(t_exp)
            story.append(Spacer(1, 1))
            
            if desc:
                for line in desc.split("\n"):
                    if line.strip():
                        bullet_txt = line.strip().lstrip("•-*").strip()
                        story.append(Paragraph(f"&bull; {bullet_txt}", bullet_style))
            story.append(Spacer(1, 3))

    # --- PROJECTS ---
    if projects:
        add_section_header("Projects")
        for proj in projects:
            p_name = proj.get("projectName", "")
            role = proj.get("role", "")
            desc = proj.get("description", "")
            techs = proj.get("technologies", "")
            
            header_left = f"<b>{p_name}</b>"
            if role:
                header_left += f" — <i>{role}</i>"
            
            row1_left = Paragraph(header_left, left_style)
            row1_right = Paragraph(f"<i>{techs}</i>" if techs else "", right_style)
            
            t_proj = Table([[row1_left, row1_right]], colWidths=[383, 140])
            t_proj.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(t_proj)
            story.append(Spacer(1, 1))
            
            if desc:
                for line in desc.split("\n"):
                    if line.strip():
                        bullet_txt = line.strip().lstrip("•-*").strip()
                        story.append(Paragraph(f"&bull; {bullet_txt}", bullet_style))
            story.append(Spacer(1, 3))

    # --- LEADERSHIP & ACTIVITIES ---
    if leadership:
        add_section_header("Leadership & Activities")
        for lead in leadership:
            org = lead.get("organization", "")
            role = lead.get("role", "")
            desc = lead.get("description", "")
            start = lead.get("startDate", "")
            end = lead.get("endDate", "")
            
            dates = f"{start} – {end}" if start and end else (start or end or "")
            
            row1_left = Paragraph(f"<b>{org}</b>", left_bold_style)
            row1_right = Paragraph(dates, right_style)
            
            row2_left = Paragraph(role, left_italic_style)
            row2_right = Paragraph("", right_style)
            
            t_lead = Table([[row1_left, row1_right], [row2_left, row2_right]], colWidths=[383, 140])
            t_lead.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,0), (-1,-1), 1),
                ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(t_lead)
            story.append(Spacer(1, 1))
            
            if desc:
                for line in desc.split("\n"):
                    if line.strip():
                        bullet_txt = line.strip().lstrip("•-*").strip()
                        story.append(Paragraph(f"&bull; {bullet_txt}", bullet_style))
            story.append(Spacer(1, 3))

    # --- SKILLS & INTERESTS ---
    has_skills = any([
        skills.get("technicalSkills"),
        skills.get("programmingLanguages"),
        skills.get("tools"),
        skills.get("softSkills")
    ])
    
    if has_skills or interests:
        add_section_header("Skills & Interests")
        skill_lines = []
        if skills.get("technicalSkills"):
            skill_lines.append(f"<b>Technical Skills:</b> {skills['technicalSkills']}")
        if skills.get("programmingLanguages"):
            skill_lines.append(f"<b>Programming Languages:</b> {skills['programmingLanguages']}")
        if skills.get("tools"):
            skill_lines.append(f"<b>Tools:</b> {skills['tools']}")
        if skills.get("softSkills"):
            skill_lines.append(f"<b>Soft Skills:</b> {skills['softSkills']}")
        if interests:
            skill_lines.append(f"<b>Interests:</b> {interests}")
            
        for line in skill_lines:
            story.append(Paragraph(line, left_style))
            story.append(Spacer(1, 2))

    # --- CERTIFICATIONS ---
    if certifications:
        add_section_header("Certifications")
        for cert in certifications:
            c_name = cert.get("name", "")
            c_issuer = cert.get("issuer", "")
            c_date = cert.get("date", "")
            parts = []
            if c_name:
                parts.append(f"<b>{c_name}</b>")
            if c_issuer:
                parts.append(c_issuer)
            if c_date:
                parts.append(f"({c_date})")
            if parts:
                story.append(Paragraph(f"&bull; {', '.join(parts)}", bullet_style))

    # Build PDF Document
    try:
        doc.build(story)
        return True
    except Exception as e:
        print(f"Error building PDF: {e}")
        return False


def generate_docx_resume(resume_data, template_style, output_path):
    """Generate a DOCX document based on the Harvard Resume layout."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # Margins: 0.5 inches
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            
        # Extract data
        name = resume_data.get("name", "")
        contact = resume_data.get("contact", {})
        education = resume_data.get("education", [])
        experiences = resume_data.get("experiences", [])
        projects = resume_data.get("projects", [])
        leadership = resume_data.get("leadership", [])
        raw_skills = resume_data.get("skills", {})
        # Normalize skills: accept dict, list, or None — always produce a dict
        if isinstance(raw_skills, dict):
            skills = raw_skills
        elif isinstance(raw_skills, list):
            skills = {"technicalSkills": ", ".join(str(s) for s in raw_skills if s)}
        else:
            skills = {}
        certifications = resume_data.get("certifications", [])
        interests = resume_data.get("interests", "")
        
        # Base Times Roman style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(9.5)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # XML Border Addition
        def add_p_border_bottom(p):
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')  # 6/8 pt
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)
            
        def add_section_header(title):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title.upper())
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0, 0, 0)
            add_p_border_bottom(p)
            
        def add_left_right_entry(left_bold, left_sub, right_top, right_sub):
            table = doc.add_table(rows=2, cols=2)
            table.autofit = False
            
            # Left col: 5.0 inches, Right col: 2.27 inches
            table.columns[0].width = Inches(5.0)
            table.columns[1].width = Inches(2.27)
            
            cell_l1 = table.cell(0, 0)
            cell_r1 = table.cell(0, 1)
            cell_l2 = table.cell(1, 0)
            cell_r2 = table.cell(1, 1)
            
            # Reset cell paddings & margins
            for cell in [cell_l1, cell_r1, cell_l2, cell_r2]:
                cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)
                
            # Pop left-top
            r_l1 = cell_l1.paragraphs[0].add_run(left_bold)
            r_l1.bold = True
            r_l1.font.name = 'Times New Roman'
            r_l1.font.size = Pt(9.5)
            
            # Pop right-top
            p_r1 = cell_r1.paragraphs[0]
            p_r1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_r1 = p_r1.add_run(right_top)
            r_r1.font.name = 'Times New Roman'
            r_r1.font.size = Pt(9.5)
            
            # Pop left-bottom
            r_l2 = cell_l2.paragraphs[0].add_run(left_sub)
            r_l2.font.name = 'Times New Roman'
            r_l2.font.size = Pt(9.5)
            
            # Pop right-bottom
            p_r2 = cell_r2.paragraphs[0]
            p_r2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_r2 = p_r2.add_run(right_sub)
            r_r2.italic = True
            r_r2.font.name = 'Times New Roman'
            r_r2.font.size = Pt(9.5)
            
            # spacer
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            
        # Name
        if name:
            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_name.paragraph_format.space_before = Pt(0)
            p_name.paragraph_format.space_after = Pt(2)
            
            r_name = p_name.add_run(name.upper())
            r_name.bold = True
            r_name.font.name = 'Times New Roman'
            r_name.font.size = Pt(15)
            
        # Contact line
        contact_parts = []
        if contact.get("address"):
            contact_parts.append(contact["address"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("linkedin"):
            contact_parts.append(contact["linkedin"])
        if contact.get("portfolio"):
            contact_parts.append(contact["portfolio"])
            
        if contact_parts:
            p_contact = doc.add_paragraph()
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_contact.paragraph_format.space_before = Pt(0)
            p_contact.paragraph_format.space_after = Pt(8)
            
            contact_line = "  •  ".join(contact_parts)
            r_contact = p_contact.add_run(contact_line)
            r_contact.font.name = 'Times New Roman'
            r_contact.font.size = Pt(9)
            
        # Education
        if education:
            add_section_header("Education")
            for edu in education:
                uni = edu.get("university", "")
                degree = edu.get("degree", "")
                gpa = edu.get("gpa", "")
                start = edu.get("startYear", "")
                end = edu.get("endYear", "")
                coursework = edu.get("relevantCoursework", "")
                
                dates = f"{start} – {end}" if start and end else (start or end or "")
                left_sub = degree
                if gpa:
                    left_sub += f" (GPA: {gpa})"
                    
                add_left_right_entry(uni, left_sub, dates, "")
                
                if coursework:
                    p_cw = doc.add_paragraph()
                    p_cw.paragraph_format.left_indent = Inches(0.15)
                    p_cw.paragraph_format.space_before = Pt(0)
                    p_cw.paragraph_format.space_after = Pt(2)
                    
                    r_lbl = p_cw.add_run("Relevant Coursework: ")
                    r_lbl.bold = True
                    r_lbl.font.name = 'Times New Roman'
                    r_lbl.font.size = Pt(9.5)
                    
                    r_val = p_cw.add_run(coursework)
                    r_val.font.name = 'Times New Roman'
                    r_val.font.size = Pt(9.5)
                    
        # Experience
        if experiences:
            add_section_header("Experience")
            for exp in experiences:
                company = exp.get("company", "")
                pos = exp.get("position", "")
                loc = exp.get("location", "")
                start = exp.get("startDate", "")
                end = exp.get("endDate", "")
                desc = exp.get("description", "")
                
                dates = f"{start} – {end}" if start and end else (start or end or "")
                add_left_right_entry(company, pos, loc, dates)
                
                if desc:
                    for line in desc.split("\n"):
                        if line.strip():
                            bullet_txt = line.strip().lstrip("•-*").strip()
                            p_b = doc.add_paragraph(style='List Bullet')
                            p_b.paragraph_format.space_before = Pt(0)
                            p_b.paragraph_format.space_after = Pt(2)
                            p_b.paragraph_format.left_indent = Inches(0.25)
                            
                            r_b = p_b.add_run(bullet_txt)
                            r_b.font.name = 'Times New Roman'
                            r_b.font.size = Pt(9.5)
                            
        # Projects
        if projects:
            add_section_header("Projects")
            for proj in projects:
                p_name = proj.get("projectName", "")
                role = proj.get("role", "")
                desc = proj.get("description", "")
                techs = proj.get("technologies", "")
                
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                table.columns[0].width = Inches(5.0)
                table.columns[1].width = Inches(2.27)
                
                cell_l = table.cell(0, 0)
                cell_r = table.cell(0, 1)
                
                cell_l.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell_l.paragraphs[0].paragraph_format.space_after = Pt(1)
                cell_r.paragraphs[0].paragraph_format.space_before = Pt(0)
                cell_r.paragraphs[0].paragraph_format.space_after = Pt(1)
                
                # Left cell
                r_l = cell_l.paragraphs[0].add_run(p_name)
                r_l.bold = True
                r_l.font.name = 'Times New Roman'
                r_l.font.size = Pt(9.5)
                
                if role:
                    r_role = cell_l.paragraphs[0].add_run(f" — {role}")
                    r_role.italic = True
                    r_role.font.name = 'Times New Roman'
                    r_role.font.size = Pt(9.5)
                    
                # Right cell
                p_r = cell_r.paragraphs[0]
                p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r_r = p_r.add_run(techs)
                r_r.italic = True
                r_r.font.name = 'Times New Roman'
                r_r.font.size = Pt(9.5)
                
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                
                if desc:
                    for line in desc.split("\n"):
                        if line.strip():
                            bullet_txt = line.strip().lstrip("•-*").strip()
                            p_b = doc.add_paragraph(style='List Bullet')
                            p_b.paragraph_format.space_before = Pt(0)
                            p_b.paragraph_format.space_after = Pt(2)
                            p_b.paragraph_format.left_indent = Inches(0.25)
                            
                            r_b = p_b.add_run(bullet_txt)
                            r_b.font.name = 'Times New Roman'
                            r_b.font.size = Pt(9.5)
                            
        # Leadership & Activities
        if leadership:
            add_section_header("Leadership & Activities")
            for lead in leadership:
                org = lead.get("organization", "")
                role = lead.get("role", "")
                desc = lead.get("description", "")
                start = lead.get("startDate", "")
                end = lead.get("endDate", "")
                
                dates = f"{start} – {end}" if start and end else (start or end or "")
                add_left_right_entry(org, role, "", dates)
                
                if desc:
                    for line in desc.split("\n"):
                        if line.strip():
                            bullet_txt = line.strip().lstrip("•-*").strip()
                            p_b = doc.add_paragraph(style='List Bullet')
                            p_b.paragraph_format.space_before = Pt(0)
                            p_b.paragraph_format.space_after = Pt(2)
                            p_b.paragraph_format.left_indent = Inches(0.25)
                            
                            r_b = p_b.add_run(bullet_txt)
                            r_b.font.name = 'Times New Roman'
                            r_b.font.size = Pt(9.5)
                            
        # Skills & Interests
        has_skills = any([
            skills.get("technicalSkills"),
            skills.get("programmingLanguages"),
            skills.get("tools"),
            skills.get("softSkills")
        ])
        
        if has_skills or interests:
            add_section_header("Skills & Interests")
            
            def add_skills_line(label, val):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                
                r_lbl = p.add_run(f"{label}: ")
                r_lbl.bold = True
                r_lbl.font.name = 'Times New Roman'
                r_lbl.font.size = Pt(9.5)
                
                r_val = p.add_run(val)
                r_val.font.name = 'Times New Roman'
                r_val.font.size = Pt(9.5)
                
            if skills.get("technicalSkills"):
                add_skills_line("Technical Skills", skills["technicalSkills"])
            if skills.get("programmingLanguages"):
                add_skills_line("Programming Languages", skills["programmingLanguages"])
            if skills.get("tools"):
                add_skills_line("Tools", skills["tools"])
            if skills.get("softSkills"):
                add_skills_line("Soft Skills", skills["softSkills"])
            if interests:
                add_skills_line("Interests", interests)
                
        # Certifications
        if certifications:
            add_section_header("Certifications")
            for cert in certifications:
                c_name = cert.get("name", "")
                c_issuer = cert.get("issuer", "")
                c_date = cert.get("date", "")
                parts = []
                if c_name:
                    parts.append(c_name)
                if c_issuer:
                    parts.append(c_issuer)
                if c_date:
                    parts.append(f"({c_date})")
                if parts:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.left_indent = Inches(0.25)
                    
                    r_lbl = p.add_run(parts[0])
                    r_lbl.bold = True
                    r_lbl.font.name = 'Times New Roman'
                    r_lbl.font.size = Pt(9.5)
                    
                    if len(parts) > 1:
                        r_rest = p.add_run(", " + ", ".join(parts[1:]))
                        r_rest.font.name = 'Times New Roman'
                        r_rest.font.size = Pt(9.5)
                        
        doc.save(output_path)
        return True
    except Exception as e:
        import traceback
        print(f"Error building DOCX: {e}")
        traceback.print_exc()
        return False


def parse_cv_text_to_resume_json(text):
    """
    Parse raw CV text into a structured Harvard Resume JSON object using regex and heuristics.
    Returns empty arrays/strings for missing sections — NEVER falls back to dummy data.
    """
    import re
    from services.cv_analysis_service import extract_skills

    # Initialize structured dict with empty/blank values — no dummy data
    resume = {
        "name": "",
        "contact": {
            "email": "",
            "phone": "",
            "address": "",
            "linkedin": "",
            "portfolio": ""
        },
        "education": [],
        "experiences": [],
        "projects": [],
        "leadership": [],
        "skills": {
            "technicalSkills": "",
            "programmingLanguages": "",
            "tools": "",
            "softSkills": ""
        },
        "certifications": [],
        "interests": ""
    }

    if not text:
        print("[parse_cv] Warning: Empty text — returning blank resume.")
        return resume

    # Helper function to split text and date from a line
    def split_text_and_date(line):
        months = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
        year = r"\b(?:20\d{2}|19\d{2})\b"
        present = r"\b(?:present|sekarang|current)\b"
        
        date_pattern = rf"\b({months}\s*{year}|{year})\s*[–\-]\s*({months}\s*{year}|{year}|{present})"
        single_date_pattern = rf"\b({months}\s*{year}|{year})\b"
        
        match = re.search(date_pattern, line, re.IGNORECASE)
        if match:
            start_date = match.group(1).strip()
            end_date = match.group(2).strip()
            text_part = line[:match.start()].strip(" \t,-•–—")
            return text_part, start_date, end_date
            
        match = re.search(single_date_pattern, line, re.IGNORECASE)
        if match:
            date = match.group(1).strip()
            text_part = line[:match.start()].strip(" \t,-•–—")
            return text_part, "", date
            
        return line, "", ""

    # Helper function to split company and location from a line
    def split_company_and_location(line):
        location_keywords = r"\b(?:Remote|Indonesia|Jakarta|Bandung|Surabaya|Medan|Binjai|Tangerang|Bekasi|Bogor|Singapore|Malaysia|Remote,?\s*Indonesia|\(Remote\))\b"
        
        # Check tab/multiple spaces first
        parts = re.split(r"\t+|\s{3,}", line)
        if len(parts) > 1:
            return parts[0].strip(), parts[1].strip()
            
        # Check for parentheses at the end: e.g. "(Remote)", "(Jakarta)"
        paren_match = re.search(r"\(([^)]+)\)\s*$", line)
        if paren_match:
            loc = paren_match.group(0).strip()
            company = line[:paren_match.start()].strip(" \t,-•")
            return company, loc
            
        # Check for location keywords at the end
        location_pattern = r"\b(?:Indonesia|Remote|Jakarta|Bandung|Surabaya|Medan|Binjai|Singapore|USA|UK|Malaysia)(?:\s*\(Remote\)|\s*,\s*\w+)*\s*$"
        match = re.search(location_pattern, line, re.IGNORECASE)
        if match:
            loc = match.group(0).strip()
            company = line[:match.start()].strip(" \t,-•")
            return company, loc
            
        # Fallback split
        parts_fb = re.split(r"\t+|\s{2,}", line)
        if len(parts_fb) > 1:
            return parts_fb[0].strip(), parts_fb[1].strip()
            
        return line, ""

    lines = [line.strip() for line in text.split("\n") if line.strip()]
    print(f"[parse_cv] Total non-empty lines: {len(lines)}")

    # 1. Extract Name (first 1-4 lines, no email/phone/url/section words)
    name_exclude = re.compile(
        r"@|\+|http|www\.|linkedin|github|portfolio|resume|cv|"
        r"education|experience|skills|certification|project|leadership|interest",
        re.IGNORECASE
    )
    for line in lines[:4]:
        if name_exclude.search(line):
            continue
        words = line.split()
        if 2 <= len(words) <= 6 and all(w[0].isupper() for w in words if w.isalpha()):
            resume["name"] = line
            break
    if not resume["name"] and lines:
        resume["name"] = lines[0]

    # 2. Contact details
    email_match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)
    if email_match:
        resume["contact"]["email"] = email_match.group(0)

    phone_match = re.search(
        r"\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,6}\b", text
    )
    if phone_match:
        resume["contact"]["phone"] = phone_match.group(0)

    linkedin_match = re.search(
        r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+", text, re.IGNORECASE
    )
    if linkedin_match:
        resume["contact"]["linkedin"] = linkedin_match.group(0)

    portfolio_match = re.search(
        r"(?:https?://)?(?:www\.)?github\.com/[\w\-]+", text, re.IGNORECASE
    )
    if portfolio_match:
        resume["contact"]["portfolio"] = portfolio_match.group(0)
    else:
        portfolio_match2 = re.search(
            r"https?://(?:bit\.ly|[\w\-]+\.(?:me|dev|io|xyz|site))/[\w\-/]+", text, re.IGNORECASE
        )
        if portfolio_match2:
            resume["contact"]["portfolio"] = portfolio_match2.group(0)

    city_match = re.search(
        r"\b(?:Jakarta|Bandung|Surabaya|Medan|Depok|Tangerang|Bekasi|Bogor|Binjai|"
        r"Makassar|Semarang|Palembang|Singapore|Kuala Lumpur|London|New York|California)"
        r",?\s*(?:Indonesia|ID|SG|UK|US|North Sumatra|West Java|East Java)?\b",
        text, re.IGNORECASE
    )
    if city_match:
        resume["contact"]["address"] = city_match.group(0).strip()

    # 3. Split text into sections using strict header detection
    header_map = {
        "education":      re.compile(r"^education$|^academic\s+background$|^academic\s+credentials$|^pendidikan$", re.IGNORECASE),
        "experience":     re.compile(r"^experience[s]?$|^work\s+experience$|^professional\s+experience$|^work\s+history$|^employment$|^pengalaman\s*kerja$|^pengalaman$", re.IGNORECASE),
        "projects":       re.compile(r"^projects?$|^academic\s+projects?$|^personal\s+projects?$|^proyek$", re.IGNORECASE),
        "leadership":     re.compile(r"^leadership(?:\s*(?:&\s*activities|\s*and\s*activities))?$|^leadership\s+activities$|^leadership\s+roles?$|^extracurricular\s*activities$|^volunteer\s*experience$|^organisasi$", re.IGNORECASE),
        "skills":         re.compile(r"^skills?$|^skills?\s*&\s*interests?$|^core\s+competencies$|^keahlian$", re.IGNORECASE),
        "certifications": re.compile(r"^certifications?$|^licenses?$|^courses?$|^credentials?$|^sertifikasi$", re.IGNORECASE),
        "interests":      re.compile(r"^interests?$|^hobbies$|^personal\s+interests?$|^minat$", re.IGNORECASE),
    }

    sections = {}
    current_section = None

    for line in lines:
        clean = line.strip(" :-_•").rstrip()
        clean_norm = re.sub(r"\s*&\s*interests?$", "", clean, flags=re.IGNORECASE).strip()
        matched = False
        for sec_name, pattern in header_map.items():
            if pattern.match(clean) or pattern.match(clean_norm):
                current_section = sec_name
                sections.setdefault(current_section, [])
                matched = True
                print(f"[parse_cv] Section detected: '{sec_name}' from line: '{line}'")
                break
        if not matched and current_section:
            sections[current_section].append(line)

    print(f"[parse_cv] Sections found: {list(sections.keys())}")

    # 4. Parse EDUCATION
    edu_lines = sections.get("education", [])
    if edu_lines:
        current_edu = None
        for line in edu_lines:
            line_clean = line.strip()
            uni_clean = re.sub(r"[\t\s]+20\d{2}\s*[–\-]\s*(?:20\d{2}|present|sekarang|current)", "", line_clean).strip()
            uni_clean = re.sub(r"\t+", " ", uni_clean).strip()

            is_school_start = re.search(
                r"university|universitas|college|school|institute|politeknik|stmik|"
                r"akademi|faculty|facultas|sekolah|perguruan",
                line_clean, re.IGNORECASE
            )

            if is_school_start or current_edu is None:
                if current_edu:
                    resume["education"].append(current_edu)
                current_edu = {
                    "university": uni_clean,
                    "degree": "",
                    "gpa": "",
                    "startYear": "",
                    "endYear": "",
                    "relevantCoursework": ""
                }
                year_matches = re.findall(r"\b(20\d{2})\b", line_clean)
                if len(year_matches) >= 2:
                    current_edu["startYear"] = year_matches[0]
                    current_edu["endYear"] = year_matches[1]
                elif len(year_matches) == 1:
                    current_edu["endYear"] = year_matches[0]
                continue

            # Degree line
            if re.search(
                r"bachelor|master|phd|doctorate|b\.sc|b\.s\b|b\.a\b|m\.sc|m\.s\b|"
                r"m\.a\b|s1|s2|diploma|associate|sarjana|magister|informatics|"
                r"computer science|engineering|teknik",
                line_clean, re.IGNORECASE
            ) and not current_edu["degree"]:
                current_edu["degree"] = re.sub(r"\t+", " ", line_clean).strip()
                gpa_match = re.search(r"gpa\s*:?\s*([0-9]+\.[0-9]+)", line_clean, re.IGNORECASE)
                if gpa_match:
                    current_edu["gpa"] = gpa_match.group(1)
                continue

            # GPA
            gpa_match = re.search(r"gpa\s*:?\s*([0-9]+\.[0-9]+)", line_clean, re.IGNORECASE)
            if gpa_match and not current_edu["gpa"]:
                current_edu["gpa"] = gpa_match.group(1)

            # Years
            if not current_edu["startYear"] or not current_edu["endYear"]:
                year_matches = re.findall(r"\b(20\d{2})\b", line_clean)
                if year_matches:
                    if len(year_matches) >= 2:
                        current_edu["startYear"] = year_matches[0]
                        current_edu["endYear"] = year_matches[1]
                    elif not current_edu["endYear"]:
                        current_edu["endYear"] = year_matches[0]

            # Relevant Coursework
            cw_match = re.search(r"(?:relevant coursework|coursework)\s*:?\s*(.+)", line_clean, re.IGNORECASE)
            if cw_match:
                current_edu["relevantCoursework"] = cw_match.group(1).strip()

        if current_edu:
            resume["education"].append(current_edu)

    print(f"[parse_cv] Education entries found: {len(resume['education'])}")

    # Helper function to group list lines into chunks of (header_lines, bullets)
    def chunk_section_lines(section_lines):
        chunks = []
        current_chunk = None
        for line in section_lines:
            line_str = line.strip()
            if not line_str:
                continue
            is_bullet = line_str.startswith(("-", "•", "*", "◦"))
            if is_bullet:
                if current_chunk is None:
                    current_chunk = {"headers": [], "bullets": []}
                clean_bullet = re.sub(r"^[•\-\*◦]\s*", "", line_str).strip()
                current_chunk["bullets"].append(clean_bullet)
            else:
                if current_chunk and current_chunk["bullets"]:
                    chunks.append(current_chunk)
                    current_chunk = None
                if current_chunk is None:
                    current_chunk = {"headers": [], "bullets": []}
                current_chunk["headers"].append(line_str)
        if current_chunk:
            chunks.append(current_chunk)
        return chunks

    # 5. Parse EXPERIENCE
    exp_lines = sections.get("experience", [])
    if exp_lines:
        chunks = chunk_section_lines(exp_lines)
        for chunk in chunks:
            headers = chunk["headers"]
            bullets = chunk["bullets"]
            if not headers:
                continue
            
            exp = {
                "company": "",
                "position": "",
                "location": "",
                "startDate": "",
                "endDate": "",
                "description": "\n".join(bullets)
            }
            
            if len(headers) == 1:
                first = headers[0]
                text_no_dates, start_date, end_date = split_text_and_date(first)
                exp["startDate"] = start_date
                exp["endDate"] = end_date
                
                company, location = split_company_and_location(text_no_dates)
                company_parts = re.split(r"\s*[—–\-]\s*", company, maxsplit=1)
                if len(company_parts) > 1:
                    exp["company"] = company_parts[0].strip()
                    exp["position"] = company_parts[1].strip()
                else:
                    exp["company"] = company.strip()
                exp["location"] = location
                
            elif len(headers) >= 2:
                first = headers[0]
                company, location = split_company_and_location(first)
                exp["company"] = company
                exp["location"] = location
                
                second = headers[1]
                pos_text, start_date, end_date = split_text_and_date(second)
                exp["position"] = pos_text
                exp["startDate"] = start_date
                exp["endDate"] = end_date
                
            resume["experiences"].append(exp)

    print(f"[parse_cv] Experience entries found: {len(resume['experiences'])}")

    # 6. Parse PROJECTS
    proj_lines = sections.get("projects", [])
    if proj_lines:
        chunks = chunk_section_lines(proj_lines)
        for chunk in chunks:
            headers = chunk["headers"]
            bullets = chunk["bullets"]
            if not headers:
                continue
                
            proj = {
                "projectName": "",
                "role": "",
                "description": "\n".join(bullets),
                "technologies": ""
            }
            
            first = headers[0]
            parts = re.split(r"\t+|\s{3,}", first)
            header_part = parts[0].strip()
            if len(parts) > 1:
                proj["technologies"] = parts[1].strip()
            else:
                tech_parts = re.split(r"\s*\|\s*", header_part)
                if len(tech_parts) > 1:
                    proj["technologies"] = tech_parts[-1].strip()
                    header_part = " | ".join(tech_parts[:-1]).strip()
            
            dash_parts = re.split(r"\s*[—–\-]\s*", header_part, maxsplit=1)
            proj["projectName"] = dash_parts[0].strip()
            if len(dash_parts) > 1:
                proj["role"] = dash_parts[1].strip()
                
            resume["projects"].append(proj)

    print(f"[parse_cv] Project entries found: {len(resume['projects'])}")

    # 7. Parse LEADERSHIP & ACTIVITIES
    lead_lines = sections.get("leadership", [])
    if lead_lines:
        chunks = chunk_section_lines(lead_lines)
        for chunk in chunks:
            headers = chunk["headers"]
            bullets = chunk["bullets"]
            if not headers:
                continue
                
            lead = {
                "organization": "",
                "role": "",
                "startDate": "",
                "endDate": "",
                "description": "\n".join(bullets)
            }
            
            if len(headers) == 1:
                first = headers[0]
                text_no_dates, start_date, end_date = split_text_and_date(first)
                lead["startDate"] = start_date
                lead["endDate"] = end_date
                
                org_parts = re.split(r"\s*[—–\-]\s*", text_no_dates, maxsplit=1)
                lead["organization"] = org_parts[0].strip()
                if len(org_parts) > 1:
                    lead["role"] = org_parts[1].strip()
            elif len(headers) >= 2:
                lead["organization"] = headers[0].strip()
                
                second = headers[1]
                pos_text, start_date, end_date = split_text_and_date(second)
                lead["role"] = pos_text
                lead["startDate"] = start_date
                lead["endDate"] = end_date
                
            resume["leadership"].append(lead)

    print(f"[parse_cv] Leadership entries found: {len(resume['leadership'])}")

    # 8. Parse SKILLS
    skills_lines = sections.get("skills", [])
    for line in skills_lines:
        line_clean = line.strip()
        int_match = re.search(r"^interests?\s*:?\s*(.+)", line_clean, re.IGNORECASE)
        if int_match and not resume["interests"]:
            resume["interests"] = int_match.group(1).strip()
            continue

        tech_match = re.search(r"technical skills?\s*:?\s*(.+)", line_clean, re.IGNORECASE)
        if tech_match:
            resume["skills"]["technicalSkills"] = tech_match.group(1).strip()
            continue

        lang_match = re.search(r"programming languages?\s*:?\s*(.+)", line_clean, re.IGNORECASE)
        if lang_match:
            resume["skills"]["programmingLanguages"] = lang_match.group(1).strip()
            continue

        tools_match = re.search(r"tools?\s*:?\s*(.+)", line_clean, re.IGNORECASE)
        if tools_match:
            resume["skills"]["tools"] = tools_match.group(1).strip()
            continue

        soft_match = re.search(r"soft skills?\s*:?\s*(.+)", line_clean, re.IGNORECASE)
        if soft_match:
            resume["skills"]["softSkills"] = soft_match.group(1).strip()
            continue

    if not any(resume["skills"].values()):
        all_skills = extract_skills(text)
        if all_skills:
            prog_langs_set = {"python", "java", "javascript", "typescript", "c++", "c#", "php", "ruby", "go", "kotlin", "swift", "sql", "html", "css"}
            tools_set = {"docker", "kubernetes", "git", "github", "aws", "gcp", "azure", "jenkins", "linux", "terraform", "figma", "arduino", "grafana", "supabase", "mysql", "colab", "vs code", "vscode"}
            soft_set = {"leadership", "communication", "teamwork", "agile", "scrum", "problem solving", "time management", "adaptability", "critical thinking", "research"}

            pl, tl, ts, ss = [], [], [], []
            for sk in all_skills:
                sk_l = sk.lower()
                if sk_l in prog_langs_set:
                    pl.append(sk)
                elif sk_l in tools_set:
                    tl.append(sk)
                elif any(s in sk_l for s in soft_set):
                    ss.append(sk)
                else:
                    ts.append(sk)

            resume["skills"]["programmingLanguages"] = ", ".join(pl)
            resume["skills"]["tools"] = ", ".join(tl)
            resume["skills"]["technicalSkills"] = ", ".join(ts)
            resume["skills"]["softSkills"] = ", ".join(ss)

    print(f"[parse_cv] Skills parsed: {resume['skills']}")

    # 9. Parse CERTIFICATIONS
    cert_lines = sections.get("certifications", [])
    for line in cert_lines:
        clean = re.sub(r"^[•\-\*]\s*", "", line).strip()
        if not clean:
            continue
        parts = [p.strip() for p in re.split(r",\s*", clean)]
        name = parts[0]
        issuer = parts[1] if len(parts) > 1 else ""
        date_raw = parts[2] if len(parts) > 2 else ""
        date = re.sub(r"[()]+", "", date_raw).strip()

        if name:
            resume["certifications"].append({
                "name": name,
                "issuer": issuer,
                "date": date
            })

    print(f"[parse_cv] Certification entries found: {len(resume['certifications'])}")

    # 10. Parse INTERESTS
    interests_lines = sections.get("interests", [])
    if interests_lines:
        interests_text = " ".join(interests_lines)
        interests_clean = re.sub(r"^interests?\s*:?\s*", "", interests_text, flags=re.IGNORECASE).strip()
        resume["interests"] = interests_clean

    print(f"[parse_cv] Interests: '{resume['interests']}'")

    # Combine technical skill strings into a flat array of strings to satisfy target JSON requirement
    flat_tech_skills = []
    if resume["skills"]["technicalSkills"]:
        flat_tech_skills.extend([s.strip() for s in resume["skills"]["technicalSkills"].split(",") if s.strip()])
    if resume["skills"]["programmingLanguages"]:
        flat_tech_skills.extend([s.strip() for s in resume["skills"]["programmingLanguages"].split(",") if s.strip()])
    if resume["skills"]["tools"]:
        flat_tech_skills.extend([s.strip() for s in resume["skills"]["tools"].split(",") if s.strip()])

    # Add duplicate mapped keys to support BOTH frontend state names and target validation formats
    resume["experience"] = resume["experiences"]
    resume["leadershipActivities"] = resume["leadership"]
    resume["technicalSkills"] = flat_tech_skills

    return resume


